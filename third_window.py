import tkinter as tk
from PIL import ImageTk, Image
import tkinter.filedialog
import cv2
import numpy as np
from array import *
import math
import random



window = tk.Tk() 
window.title(u"523 Recognition Of Successful Shots On Targets (ROSOT)")
window.geometry("680x500")


with open('threshold.txt') as f:
	data = f.readline()
	threshold = int(data.strip())
f.close()
	
with open("apotelesmata_xw.txt") as file_in:
		lines = []
		for line in file_in:
			lines.append(line)
file_in.close()
given_h=0
given_w = 0
given_h = int(lines[0])
given_w = int(lines[1])
	


def select_first_image():
	file1=tkinter.filedialog.askopenfilename()
	paths.append(file1)
	
def select_second_image():
	file2=tkinter.filedialog.askopenfilename()
	paths.append(file2)

def pointInEllipse(x,y,xp,yp,d,D,angle):
    #tests if a point[xp,yp] is within
    #boundaries defined by the ellipse
    #of center[x,y], diameter d D, and tilted at angle

    cosa=math.cos(angle)
    sina=math.sin(angle)
    dd=d/2*d/2
    DD=D/2*D/2

    a =(cosa*(xp-x)+sina*(yp-y))**2
    b =(sina*(xp-x)-cosa*(yp-y))**2
    ellipse=(a/dd)+(b/DD)

    if ellipse <= 1:
        return True
    else:
        return False
		


counter_seiras_volis=1
paths=[]

def efarmogi():
	
	f=open("onomata.txt","r",encoding='utf-8')

	array=f.readlines()


	image = cv2.imread(str(paths[0]))

	
	
	gray=cv2.cvtColor(image,cv2.COLOR_BGR2GRAY)
	
	thresh=threshold
	ret,thresh_img=cv2.threshold(gray,thresh,255,cv2.THRESH_BINARY)
	
	new_img=cv2.resize(thresh_img,(1000,1000))
	cv2.imshow("thresh",new_img)
	cv2.waitKey(0)
	contours,hierarchy=cv2.findContours(thresh_img,cv2.RETR_TREE,cv2.CHAIN_APPROX_SIMPLE)
	
	img_contours=np.zeros(image.shape)
	cv2.drawContours(img_contours,contours,-1,(0,255,0),3)
	new_img=cv2.resize(img_contours,(1000,1000))
	cv2.imshow("thresh",img_contours)
	cv2.waitKey(0)
	
	#sorted_contours= sorted(contours,key=lambda ctr:cv2.boundingRect(ctr))
	sorted_contours= sorted(contours,key=lambda ctr:cv2.contourArea(ctr), reverse = True)
	
	image_array =[]
	overlap_x =[]
	overlap_w =[]
	overlap_idx = 0
	see_f=0
	for c in sorted_contours:
	    #koftis gia > apo stoxous user
		
		x,y,w,h = cv2.boundingRect(c)	

		width_pc= given_w * 0.4
		height_pc = given_h * 0.4
		
		overlap = True 
		if w>given_w-width_pc and h>given_h-height_pc and w<given_w+width_pc and h<given_h+height_pc :	 #(prev_x >=x or x>= prev_xw or prev_x >= xw or xw>= prev_xw):
			if(overlap_idx == 0):
				
				overlap_idx = 0  #dummy command
			else:
				
				for k in range(len(overlap_x)):
					if not((overlap_x[k]>x and x +w >overlap_x[k] + overlap_w[k]) or ( x+w >overlap_x[k] and x< overlap_x[k]+overlap_w[k]) or (overlap_x[k] == x and overlap_x[k]+overlap_w[k] == x+w)):
						
						overlap = False
					else:
						overlap = True
						break
				if overlap == True: continue
			new_img=image[y:y+h,x:x+w]
			#cv2.imshow("Canny Edge",new_img)
			sc2 =cv2.resize(new_img,(700,700))
			#cv2.imshow("Targets detected",sc2)
			cv2.waitKey(0)
			overlap_x.append(x)
			overlap_w.append(w)
			overlap_idx+=1
			image_array.append(new_img)
	#print("sintetagmens", overlap_x, overlap_w)
	
	#sort images by x coordinate
	
	sorted_image_array =[]
	sorted_image_array_idx = 0
	for j in range(len(overlap_x)):
		min_x = math.inf
		min_idx= -1
		for k in range(len(overlap_x)):
			if overlap_x[k] < min_x :
				min_x = overlap_x[k]
				min_idx= k
		overlap_x[min_idx] = math.inf
		sorted_image_array.append(image_array[min_idx])
		sorted_image_array_idx += 1
		
	for j in range(len(sorted_image_array)):
		image_array[j] = sorted_image_array[j]
		cv2.imshow("thresh",image_array[j])
		cv2.waitKey(0)
	
	
	


	megalos_kuklos1=[[5] for r in range(len(image_array))]
	megalos_kuklos2=[[5] for r in range(len(image_array))]
	mauro1=[[5] for r in range(len(image_array))]

	mikros_kuklos1=[[5] for r in range(len(image_array))]
	mikros_kuklos2=[[5] for r in range(len(image_array))]
	mauro2=[[5] for r in range(len(image_array))]

	
	circles=[]
	
	for i in range(len(image_array)):

		gray4=cv2.cvtColor(image_array[i],cv2.COLOR_BGR2GRAY)
		#sp =cv2.resize(image_array[0],(700,700))
		#cv2.imshow("prwtos",sp)
		#cv2.waitKey(0)
		thresh=threshold
		# gray4 = cv2.medianBlur(gray4,3,3)
		ret4,thresh_img4=cv2.threshold(gray4,thresh,255,cv2.THRESH_BINARY)
		#cv2.imshow("palia",thresh_img4)
		#cv2.waitKey(0)


		contours4,hierarchy4=cv2.findContours(thresh_img4,cv2.RETR_TREE,cv2.CHAIN_APPROX_SIMPLE)

		img_contours4=np.zeros(image_array[i].shape)
		cv2.drawContours(img_contours4,contours4,-1,(0,255,0),3)
		
		sorted_contours= sorted(contours4,key=lambda ctr:cv2.contourArea(ctr), reverse = True)
		
		
		#for k in range(len(sorted_contours)):
		ellipse1 = cv2.fitEllipse(sorted_contours[0])
		img_contours4=image_array[i].copy()
		cv2.ellipse(img_contours4,ellipse1,(0,0,255),2)
		cv2.imshow("prwtos",img_contours4)
		#print(ellipse)
		cv2.waitKey(0)
		
		ellipse2 = cv2.fitEllipse(sorted_contours[1])
			
		
		print(ellipse1)
		print(ellipse2)
		sfaires_mesa=0
		sfaires_mesa2=0
	
		for c in contours4:
			(x1, y1, w1, h1) = cv2.boundingRect(c)
			center_point_x = x1+w1//2
			center_point_y = y1+h1//2
			
			current_bullet = image_array[i]
			if cv2.contourArea(c) > 1 and  cv2.contourArea(c)<100:
				test_point1 = pointInEllipse(ellipse1[0][0],ellipse1[0][1],center_point_x,center_point_y,ellipse1[1][0],ellipse1[1][1],ellipse1[2])
				if test_point1 == True:
					sfaires_mesa+=1
					megalos_kuklos2[i][0]=sfaires_mesa
					cv2.drawContours(current_bullet,c,-1,(0,255,0),3)
					cv2.imshow("Current_bullet megalos", current_bullet)
					cv2.waitKey(0)
					
		for c in contours4:
			(x1, y1, w1, h1) = cv2.boundingRect(c)
			center_point_x = x1+w1//2
			center_point_y = y1+h1//2	
			current_bullet = image_array[i]
			if cv2.contourArea(c) > 1 and  cv2.contourArea(c)<100:
				test_point2 = pointInEllipse(ellipse2[0][0],ellipse2[0][1],center_point_x,center_point_y,ellipse2[1][0],ellipse2[1][1],ellipse2[2])
				if test_point2 == True:
					sfaires_mesa2+=1
					mikros_kuklos2[i][0]=sfaires_mesa2
					cv2.drawContours(current_bullet,c,-1,(255,0,0),3)
					cv2.imshow("Current_bullet mikros", current_bullet)
					cv2.waitKey(0)
		
		megalos_kuklos2[i][0] = megalos_kuklos2[i][0] - mikros_kuklos2[i][0]
		print("PHOTO 1: mikros kuklos,megalos kuklos",mikros_kuklos2[i][0],megalos_kuklos2[i][0])		
				
		#for k in range(len(sorted_contours)):

		canny = cv2.Canny(gray4,thresh,thresh)
		canny = cv2.medianBlur(canny,3,3)
		cv2.imshow("canny",canny)
		cv2.waitKey(0)
		
		
		contoursc,hierarchy4=cv2.findContours(canny,cv2.RETR_TREE,cv2.CHAIN_APPROX_SIMPLE)

		canny_contours=np.zeros(image_array[i].shape)
		cv2.drawContours(canny_contours,contoursc,-1,(255,0,0),3)
		cv2.imshow("Canny contours",canny_contours)
		cv2.waitKey(0)
					
		mauro = 0
		for c in contoursc:
			(x1, y1, w1, h1) = cv2.boundingRect(c)
			center_point_x = x1+w1//2
			center_point_y = y1+h1//2	
			current_bullet = image_array[i]
			if cv2.contourArea(c) > 1 and  cv2.contourArea(c)<100:
				test_point2 = pointInEllipse(ellipse2[0][0],ellipse2[0][1],center_point_x,center_point_y,ellipse2[1][0],ellipse2[1][1],ellipse2[2])
				if test_point2 == True:
					mauro+=1
					mauro2[i][0]=mauro
					cv2.drawContours(current_bullet,c,-1,(255,0,0),3)
					cv2.imshow("Current_bullet mauros", current_bullet)
					cv2.waitKey(0)
		
		mauro2[i][0] = mauro2[i][0] - mikros_kuklos2[i][0]
		print("mauroooo", mauro2[i][0])
		
		print("sunolo =", mauro2[i][0]+ mikros_kuklos2[i][0] + megalos_kuklos2[i][0])
		
		
		#cv2.imshow("prwtos",img_contours4)
		#cv2.waitKey(0)
		
		#gray3 = cv2.cvtColor(image_array[i], cv2.COLOR_BGR2GRAY)
		#gray_blur3 = cv2.GaussianBlur(gray3,3,3)
		
		#gray3 = cv2.Canny(gray3 ,30, 150, L2gradient = True)
		#cv2.imshow("hough",gray3)
		#cv2.waitKey(0)
		
		
		#contours4,hierarchy4=cv2.findContours(gray3,cv2.RETR_TREE,cv2.CHAIN_APPROX_SIMPLE)

		#img_contours4=np.zeros(image_array[i].shape)
		#cv2.drawContours(img_contours4,contours4,-1,(0,255,0),3)
		
		#cv2.imshow("countour canny",img_contours4)
		#cv2.waitKey(0)
	
		
	
	
	
		
		
		
		


	#########EPEKSERGASIA GIA THN DEUTERI EIKONA######################################

	#reading image
	image2 = cv2.imread(paths[1])
	
	
	gray2=cv2.cvtColor(image2,cv2.COLOR_BGR2GRAY)
	
	thresh=threshold
	ret2,thresh_img2=cv2.threshold(gray2,thresh,255,cv2.THRESH_BINARY)
	


	contours2,hierarchy2=cv2.findContours(thresh_img2,cv2.RETR_TREE,cv2.CHAIN_APPROX_SIMPLE)

	img_contours2=np.zeros(image2.shape)
	
	sorted_contours2= sorted(contours,key=lambda ctr:cv2.contourArea(ctr), reverse = True)
	#sorted_contours2= sorted(contours2,key=lambda ctr:cv2.boundingRect(ctr))

	
	image_array2 =[]
	overlap_x =[]
	overlap_w =[]
	overlap_idx = 0
	for c in sorted_contours2:

		x,y,w,h = cv2.boundingRect(c)
		width_pc= given_w * 0.4
		height_pc = given_h * 0.4
		if w>given_w-width_pc and h>given_h-height_pc and w<given_w+width_pc and h<given_h+height_pc :
			if overlap_idx==0:
				overlap_idx = 0  #dummy command
				
			else:
					
				for k in range(len(overlap_x)):
					if not((overlap_x[k]>x and x +w >overlap_x[k] + overlap_w[k]) or ( x+w >overlap_x[k] and x< overlap_x[k]+overlap_w[k]) or (overlap_x[k] == x and overlap_x[k]+overlap_w[k] == x+w)):
						
						overlap = False
					else:
						overlap = True
						break
				if overlap == True: continue
			new_img=image2[y:y+h,x:x+w]
			#cv2.imshow("Canny Edge",new_img)
			sc2 =cv2.resize(new_img,(700,700))
			#cv2.imshow("Targets detected",sc2)
			#cv2.waitKey(0)
			overlap_x.append(x)
			overlap_w.append(w)
			overlap_idx+=1
			image_array2.append(new_img)
	print("PHOTO 1 suntetagmenes",overlap_x, overlap_w)
	
	
	sorted_image_array2 =[]
	sorted_image_array2_idx = 0
	for j in range(len(overlap_x)):
		min_x = math.inf
		min_idx= -1
		for k in range(len(overlap_x)):
			if overlap_x[k] < min_x :
				min_x = overlap_x[k]
				min_idx= k
		overlap_x[min_idx] = math.inf
		sorted_image_array2.append(image_array2[min_idx])
		sorted_image_array2_idx += 1
		
	for j in range(len(sorted_image_array2)):
		image_array2[j] = sorted_image_array2[j]

	circles=[]
	
	for i in range(len(image_array2)):

		gray4=cv2.cvtColor(image_array2[i],cv2.COLOR_BGR2GRAY)
		thresh=threshold
		ret4,thresh_img4=cv2.threshold(gray4,thresh,255,cv2.THRESH_BINARY)
		
		contours4,hierarchy4=cv2.findContours(thresh_img4,cv2.RETR_TREE,cv2.CHAIN_APPROX_SIMPLE)

		img_contours4=np.zeros(image_array2[i].shape)
		cv2.drawContours(img_contours4,contours4,-1,(0,255,0),3)
		
		sorted_contours= sorted(contours4,key=lambda ctr:cv2.contourArea(ctr), reverse = True)
		
		img_contours4=np.zeros(image_array2[i].shape)
		
		(x,y),radius = cv2.minEnclosingCircle(sorted_contours[0])
		center = (int(x),int(y))
		radius = int(radius)
		cv2.circle(image_array2[i],center,radius,(255,0,0),2)
		
		
		#cv2.imshow("contours ana megethos",image_array2[i])#########################################
		cv2.waitKey(0)
		
	
		circles.append((int(x),int(y),int(radius)))
		
				
		if circles[i] is not None:
		
			#circles[i] = np.round(circles[i][0:2]).astype("int")
		
			#for (x, y, r) in circles[i]:
			for k in range(len(circles)):
				x = circles[k][0]
				y = circles[k][1]
				r = circles[k][2]
				sfaires_mesa=0
				sfaires_mesa2=0
				for c in contours4:
				
					(x1, y1, w1, h1) = cv2.boundingRect(c)
				
					#if(r<200):
					#	r=242
					
					cv2.circle(image_array2[i], (x, y), r, (0, 255, 0), 4)
					#cv2.imshow("hough",image_array[i])
					#cv2.waitKey(0)
					
					apostasi=math.sqrt( ((x1 - x)**2) + ((y1 - y)**2) )
						
					if(apostasi<=(r/2) and cv2.contourArea(c) > 1 and  cv2.contourArea(c)<100):
						
						sfaires_mesa+=1
						mikros_kuklos2[i][0]=sfaires_mesa
						
					if(apostasi>(r/2) and apostasi<=r and cv2.contourArea(c) > 1 and  cv2.contourArea(c)<100):
						
						sfaires_mesa2+=1	
						
						megalos_kuklos2[i][0]=sfaires_mesa2
						
			print("PHOTO 2: mikros kuklos,megalos kuklos",sfaires_mesa,sfaires_mesa2)
			cv2.imshow("hough",image_array2[i])
			cv2.waitKey(0)
			
			#################################
	
					
	
	stoxos = []
	onomata = []
	points = [] 
	for k in range (len(image_array2)):
		points.append(abs(mikros_kuklos2[k][0]-mikros_kuklos1[k][0])*5 + abs(megalos_kuklos2[k][0]-megalos_kuklos1[k][0])*4)
		stoxos.append(abs(mikros_kuklos2[k][0]-mikros_kuklos1[k][0]) + abs(megalos_kuklos2[k][0]-megalos_kuklos1[k][0]))
		onomata.append(array[k].strip("\n")+": (στον στόχο)"+str(stoxos[k]) + " (πόντοι)" + str(points[k]) +"\n")

	paths.clear()
	
	apotelesma = ""
	for k in range (len(image_array2)):
		apotelesma = apotelesma + onomata[k] + "\n"
		
	apotelesmata_volis.insert(tk.END,apotelesma)

	file = open("telika_apotelesmata.txt", "w",encoding='utf-8')

	file.write(apotelesma)
	file.close()
	
from tkinter import *	
def temp():
	apotelesmata_volis.insert(tk.END,"Εμφάνιση Αποτελεσμάτων\nΠαρακαλώ Περιμένετε!")
	efarmogi(event)
def test():
	#window.destroy()

	close_window()

def close_window():
	window1 = Toplevel()
	window1.title(u"523 Recognition Of Successful Shots On Targets (ROSOT)")
	#window.resizable(800,600)
	window1.geometry("680x360")
	f=open("informations.txt","r",encoding='utf-8')


	seires=f.readlines()
	temp3=seires[3].split(":")
	

	array=[]
	array2=[]
	counter_lines=0
	for i in range(int(temp3[1])):
		counter_lines+=1
		temp_str="Όνομα Σκοπευτή "+str(counter_lines)+"ου Στόχου"
		temp=tk.Label(window1, text=temp_str).grid(row=(counter_lines+1))
		temp2 = tk.Entry(window1,width=50)
		temp2.grid(row=(counter_lines+1),column=1)
		array.append(temp)
		array2.append(temp2)


	

	button1=tk.Button(window1, text="Eπόμενο")
	button1.place(x=580, y=300)

	information_array=[]

	def close_second_window():
		file = open("onomata.txt", "w+",encoding='utf-8')
		#file.write(str((e1.get()))+"\n")
		for i in range(int(temp3[1])):
			temp=array2[i].get()
			file.write(str((temp))+"\n")
		
	
		file.close()
		
		
		
		window1.withdraw()
		#window.destroy()
		#create_new_window('<Button-1>')
	button1=tk.Button(window1, text="Eπόμενο",command=close_second_window)
	button1.place(x=580, y=300) 





def word():
	from docx import Document
	from docx.shared import Inches

	document = Document()

	f1=open("onomata.txt","r",encoding='utf-8')
	f2=open("apotelesmata.txt","r",encoding='utf-8')
	f3=open("informations.txt","r",encoding='utf-8')
	f4=open("telika_apotelesmata.txt","r",encoding='utf-8')

	array4=f4.readlines()
	onomata=[]
	sfaires=[]
	scores=[]
	for i in range(len(array4)):
		onoma=array4[i].split(":")
		sfaires1=onoma[0]
		temp=onoma[1].split("/")
		sfaires_mesa=temp[0]
		skor=temp[1]
		onomata.append(sfaires1)
		sfaires.append(sfaires_mesa)
		scores.append(skor)


	array=f3.readlines()
	print(array)
	temp5=array[1].split(":")
	document.add_heading(array[2], 0)
	document.add_heading('Αποτελέσματα Βολής της'+str(temp5[1]), 0)
	for i in range(len(array)):
		if(i!=1 and i!=2):
			print(array[i])
			document.add_paragraph(array[i])
		
	#p = document.add_paragraph(array)
	document.add_paragraph("")
	document.add_paragraph("")
	document.add_paragraph("")
	document.add_paragraph("")
	document.add_paragraph("")
	document.add_heading('Αποτελέσματα Σκοπευτών', level=1)
	document.add_paragraph('', style='Intense Quote')


	#document.add_picture('monty-truth.png', width=Inches(1.25))

	records = (
		(3, '101', 'Spam'),
		(7, '422', 'Eggs'),
		(4, '631', 'Spam, spam, eggs, and spam')
	)
	
	table = document.add_table(rows=1, cols=3)
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = 'Σκοπευτής'
	hdr_cells[1].text = 'Σφαίρες στον Στόχο'
	hdr_cells[2].text = 'Επίδοση'
	row_cells=0





	print(scores)
	for i in range(len(onomata)):
		row_cells = table.add_row().cells
		row_cells[0].text = str(onomata[i])
		row_cells[1].text = str(sfaires[i])
		row_cells[2].text = str(scores[i])
	#for qty in sfaires:
	#	row_cells[1].text = str(qty)

	#for qty in scores:
	#	row_cells[2].text = str(qty)

	document.add_page_break()

	document.save('Αποτελέσματα Βολής.docx')
	apotelesmata_volis.insert(tk.END,"Η Εξαγωγή του Αρχείου Πέτυχε")

button1=tk.Button(window, text="Επιλογή Πρώτης Φωτογραφίας",command= select_first_image)
button1.place(x=0, y=30,width=250)

button2=tk.Button(window, text="Επιλογή Δεύτερης Φωτογραφίας",command=select_second_image)
button2.place(x=0, y=60,width=250)

button_enarksi=tk.Button(window,text="Εμφάνιση Αποτελεσμάτων",command=efarmogi)
button_enarksi.place(x=0, y=90,width=250)

apotelesmata_volis=tk.Text(window)
apotelesmata_volis.place(x=0,y=120,width=600)
apotelesmata_volis.insert(tk.END,"Αποτελέσματα:\n")

telos=tk.Button(window,text="Συμπλήρωση Σκοπευτών",command=test)
telos.place(x=0, y=0,width=250)

word=tk.Button(window,text="Εξαγωγή σε αρχείο Word",command=word)
word.place(x=300, y=0,width=250)
#print(file)
window.mainloop()